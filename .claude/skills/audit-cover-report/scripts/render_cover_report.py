# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx", "openpyxl"]
# ///
"""Render ใบปะหน้า + หน้ารายงาน from a client CONTEXT.md by filling template tokens.

Deterministic: parses the CONTEXT Client Profile table, picks the CPA or TA
template by entity type, substitutes «TOKEN»s, and writes both .docx into the
client's WP/ dir. No free-form generation.

Usage (from repository root):
    uv run .claude/skills/audit-cover-report/scripts/render_cover_report.py "PATH/TO/6_ผลจากสกิล/<client>/CONTEXT.md" [--kind cpa|ta]

Prints JSON: chosen kind, output paths, and any required field still unresolved
(the script refuses to write if a required value is missing or still ⚠).
"""
import argparse
import json
import re
import sys
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import RGBColor

SKILL = Path(__file__).resolve().parent.parent
ASSETS = SKILL / "assets"
AUDITOR_DB = "Database ข้อมูลผู้สอบ.csv.xlsx"
FIRM_CITY = "ขอนแก่น"  # constant across all 15 sample cases (KSK office)
FIRM_OFFICE = "9/38 ถ.กลางเมือง ต.เมืองเก่า อ.เมือง จ.ขอนแก่น"  # KSK office; used when auditor has none in DB
THAI_MONTHS = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
               "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]


def fail(msg: str) -> None:
    print(json.dumps({"ok": False, "error": msg}, ensure_ascii=False))
    sys.exit(1)


def parse_context(path: Path) -> dict:
    """Extract the Client Profile table into {field: value}."""
    profile = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        m = re.match(r"\|\s*([a-z_0-9]+)\s*\|\s*(.*?)\s*\|", line)
        if m:
            profile[m.group(1)] = m.group(2)
    return profile


def parse_ddmmyyyy(raw: str) -> tuple[int, int, int] | None:
    """Parse a dd/mm/yyyy CONTEXT date field into (day, month, year) ints. Year is
    whatever calendar the field uses as-is (CONTEXT is Buddhist for period_end/sign_date
    but Gregorian for incorporation_date — callers convert as needed)."""
    m = re.match(r"(\d{1,2})/(\d{1,2})/(\d{4})", raw or "")
    if not m:
        return None
    d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
    if not 1 <= mo <= 12:
        return None
    return d, mo, y


def parse_thai_long_date(s: str) -> tuple[int, int, int] | None:
    """Reverse of thai_long_date(): parse 'D Month [พ.ศ.] YYYY' back into (day, month, year).
    Used for --liquidation-date, which is supplied as Thai long form, not dd/mm/bbbb."""
    m = re.match(r"(\d{1,2})\s+(\S+)\s+(?:พ\.ศ\.?\s*)?(\d{4})", s or "")
    if not m:
        return None
    month = m.group(2)
    if month not in THAI_MONTHS:
        return None
    return int(m.group(1)), THAI_MONTHS.index(month), int(m.group(3))


def thai_long_date(ddmmbbbb: str) -> str | None:
    parts = parse_ddmmyyyy(ddmmbbbb)
    if not parts:
        return None
    d, mo, y = parts
    return f"{d} {THAI_MONTHS[mo]} {y}"


def thai_long_date_era(ddmmbbbb: str, spaced: bool = True) -> str | None:
    """Same as thai_long_date() but with the พ.ศ. era marker inserted before the year.
    Ground truth is inconsistent about the exact spacing between contexts, so `spaced`
    picks between them:
      - cover page ("ณ วันที่ 31 ธันวาคม พ.ศ. 2568"): spaced=True
      - ปีแรก period-from-incorporation clause ("...พ.ศ.2568"): spaced=False
    NOT used for the plain report/TA STMT_DATE tokens — ground truth never puts พ.ศ. there."""
    parts = parse_ddmmyyyy(ddmmbbbb)
    if not parts:
        return None
    d, mo, y = parts
    era = "พ.ศ. " if spaced else "พ.ศ."
    return f"{d} {THAI_MONTHS[mo]} {era}{y}"


def date_suffix(ddmmbbbb: str) -> str | None:
    """ddmmyy filename suffix (e.g. '311268'), matching the convention already used by
    `1 Planning <company> ddmmyy.xlsx` and every ground-truth WP/2 หน้ารายงาน* filename."""
    parts = parse_ddmmyyyy(ddmmbbbb)
    if not parts:
        return None
    d, mo, y = parts
    return f"{d:02d}{mo:02d}{y % 100:02d}"


def detect_first_year(incorp_raw: str, prior_period_end_raw: str, period_end_raw: str) -> tuple[bool, str | None]:
    """A ปีแรก (first-year) client is one incorporated during the audited period. CONTEXT's
    incorporation_date comes straight from DBD in the Gregorian (ค.ศ.) calendar, while every
    other CONTEXT date (period_end, prior_period_end, sign_date) is Buddhist (พ.ศ.) — convert
    before comparing so this never misfires on a client that has simply always existed.
    Returns (is_first_year, incorporation date in Thai long form with unspaced พ.ศ., or None)."""
    inc = parse_ddmmyyyy(incorp_raw)
    pe = parse_ddmmyyyy(period_end_raw)
    if not inc or not pe:
        return False, None
    d, mo, y_ce = inc
    y_be = y_ce + 543
    inc_key = (y_be, mo, d)
    pe_key = (pe[2], pe[1], pe[0])
    if inc_key > pe_key:
        return False, None  # incorporated after period end — nonsensical data, don't guess
    ppe = parse_ddmmyyyy(prior_period_end_raw)
    if ppe:
        ppe_key = (ppe[2], ppe[1], ppe[0])
        if inc_key <= ppe_key:
            return False, None  # already existed as of the prior period — not first year
    return True, f"{d} {THAI_MONTHS[mo]} พ.ศ.{y_be}"


def clean(v: str) -> str | None:
    """Return the value only if resolved (no ⚠ / dash / FILL marker), with internal
    multi-spaces collapsed so DB spacing quirks don't leak into the document."""
    if not v or v.strip() in {"—", "-"} or "⚠" in v or "⟨FILL" in v:
        return None
    return re.sub(r"\s{2,}", " ", v.strip())


def sub_tokens(doc: Document, mapping: dict[str, str]) -> None:
    """Substitute tokens. Empty values are skipped so their «TOKEN» survives and is
    caught by the residual-token scan (a missing value must never render as blank)."""
    for p in doc.paragraphs:
        if "«" not in p.text or not p.runs:
            continue
        text = p.text
        for tok, val in mapping.items():
            if val:
                text = text.replace(tok, val)
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""


def drop_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def residual_tokens(path: Path) -> list[str]:
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    return sorted(set(re.findall(r"«[^»]+»", txt)))


MISSING_RED = RGBColor(0xC0, 0x00, 0x00)


def flag_unfilled(doc: Document) -> list[str]:
    """Any «TOKEN» left after sub_tokens() means a required-but-unavailable external
    fact (e.g. TA's prior_sign_date, which lives on last year's signed report, not any
    DB). Rewrite it in place as a bold red "ไม่มีข้อมูล" warning so a reviewer opening
    the .docx sees it immediately, instead of a raw brace token easy to miss. Returns
    the token names found, for the caller's own warning list."""
    found = []
    for para in doc.paragraphs:
        toks = re.findall(r"«[^»]+»", para.text)
        if not toks:
            continue
        found.extend(toks)
        text = para.text
        for tok in toks:
            text = text.replace(tok, f"[ไม่มีข้อมูล — กรุณากรอก {tok.strip('«»')}]")
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        para.runs[0].font.color.rgb = MISSING_RED
        para.runs[0].font.bold = True
    return sorted(set(found))


def auditor_extra(first_or_full: str) -> dict:
    """Look up national ID and office address from Database ผู้สอบ (needed for TA report)."""
    root = Path.cwd()
    if not (root / AUDITOR_DB).exists():
        return {}
    wb = openpyxl.load_workbook(root / AUDITOR_DB, data_only=True)
    ws = wb.active
    header = [str(c.value).strip() if c.value else "" for c in ws[1]]
    idx = {h: i for i, h in enumerate(header)}
    for row in ws.iter_rows(min_row=2, values_only=True):
        name = str(row[idx.get("ชื่อ", -1)] or "")
        if name and name in first_or_full:
            nid = re.sub(r"\D", "", str(row[idx.get("เลข บัตร ปชช", -1)] or ""))
            spaced = "  ".join(nid) if len(nid) == 13 else ""
            office = str(row[idx.get("ชื่อ ที่อยู่ สนง", -1)] or "").strip()
            return {"national_id": spaced, "office": office}
    return {}


def render_cover(kind: str, company: str, stmt_date: str, out_dir: Path) -> Path:
    doc = Document(ASSETS / f"cover_{kind}.docx")
    sub_tokens(doc, {"«COMPANY»": company, "«STMT_DATE»": stmt_date})
    out = out_dir / "ใบปะหน้างบการเงิน.docx"
    doc.save(out)
    return out


def prior_auditor_sentence(p: dict) -> str:
    """Compose the เรื่องอื่นๆ (predecessor-auditor / ISA 710) sentence deterministically
    from atomic components read out of CONTEXT. `prior_auditor_text`, if present, wins as a
    verbatim override; otherwise the sentence is assembled from opinion + sign date so the
    format never drifts. Opinion defaults to unqualified; sign-date clause is dropped if absent."""
    override = clean(p.get("prior_auditor_text") or "")
    if override:
        return override
    opinion = clean(p.get("prior_auditor_opinion") or "") or "อย่างไม่มีเงื่อนไข"
    sign_date = clean(p.get("prior_auditor_sign_date") or "")
    txt = (f"งบการเงินของ {p['company']} สำหรับปีสิ้นสุดวันที่ {p['prior_stmt_date']} "
           f"ตรวจสอบโดยผู้สอบบัญชีอื่น ซึ่งแสดงความเห็น{opinion}")
    if sign_date:
        txt += f" ตามรายงานลงวันที่ {sign_date}"
    return txt


def render_report_cpa(p: dict, out_dir: Path) -> Path:
    doc = Document(ASSETS / "report_cpa.docx")
    has_prior = (p.get("has_prior_auditor_note") or "").lower().startswith("true")
    for para in list(doc.paragraphs):
        t = para.text
        if "«IF_PRIOR»" in t:
            if has_prior:
                para.runs[0].text = t.replace("«IF_PRIOR»", "")
                for r in para.runs[1:]:
                    r.text = ""
            else:
                drop_paragraph(para)
        elif "«PRIOR_AUDITOR_TEXT»" in t:
            if has_prior:
                para.runs[0].text = prior_auditor_sentence(p)
                for r in para.runs[1:]:
                    r.text = ""
            else:
                drop_paragraph(para)
    sub_tokens(doc, {
        "«COMPANY»": p["company"], "«STMT_DATE»": p["stmt_date"],
        "«PRIOR_STMT_DATE»": p["prior_stmt_date"], "«AUDITOR_NAME»": p["auditor_name"],
        "«LICENSE»": p["license"], "«SIGN_DATE»": p["sign_date"], "«CITY»": p["city"],
        "«PERIOD_PHRASE»": p["period_phrase"],
    })
    out = out_dir / f"2 หน้ารายงาน {p['company']} {p['date_suffix']}.docx"
    doc.save(out)
    return out


def render_report_cpa_liquidation(p: dict, out_dir: Path) -> Path:
    doc = Document(ASSETS / "report_cpa_liquidation.docx")
    sub_tokens(doc, {
        "«COMPANY»": p["company"], "«LIQ_DATE»": p["liq_date"], "«PERIOD_START»": p["period_start"],
        "«AUDITOR_NAME»": p["auditor_name"], "«LICENSE»": p["license"],
        "«SIGN_DATE»": p["sign_date"], "«CITY»": p["city"],
    })
    out = out_dir / f"2 หน้ารายงานผู้สอบ {p['company']} {p['date_suffix']}.docx"
    doc.save(out)
    return out


def render_report_ta(p: dict, out_dir: Path) -> tuple[Path, list[str]]:
    doc = Document(ASSETS / "report_ta.docx")
    sub_tokens(doc, {
        "«COMPANY»": p["company"], "«STMT_DATE»": p["stmt_date"],
        "«CUR_YEAR»": p["cur_year"], "«PRIOR_YEAR»": p["prior_year"],
        "«AUDITOR_NAME»": p["auditor_name"], "«LICENSE»": p["license"],
        "«NATIONAL_ID»": p["national_id"], "«OFFICE»": p["office"],
        "«SIGN_DATE»": p["sign_date"], "«PRIOR_SIGN_DATE»": p.get("prior_sign_date", "«PRIOR_SIGN_DATE»"),
    })
    flagged = flag_unfilled(doc)
    out = out_dir / f"2 รายงานการตรวจสอบและรับรองบัญชี {p['company']}.docx"
    doc.save(out)
    return out, flagged


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("context")
    ap.add_argument("--kind", choices=["cpa", "cpa_liq", "ta"], help="override template choice")
    ap.add_argument("--prior-sign-date", help="TA only: prior-year report date (not in any DB — ask the client)")
    ap.add_argument("--office", help="TA only: override the auditor's office address line")
    ap.add_argument("--liquidation-date", help="งบเลิก only: dissolution date, Thai long form e.g. '26 ธันวาคม พ.ศ. 2568'")
    ap.add_argument("--period-start", help="งบเลิก only: period start (default 1 มกราคม พ.ศ. <liq year>)")
    args = ap.parse_args()

    ctx_path = Path(args.context)
    if not ctx_path.is_file():
        fail(f"CONTEXT.md not found: {ctx_path}")
    prof = parse_context(ctx_path)

    # Prefer the DBD full legal name; fall back to the (often abbreviated) DB name.
    company = clean(prof.get("company_legal_name")) or clean(prof.get("company_name"))
    entity = prof.get("entity_type", "")
    period_end_raw = clean(prof.get("period_end")) or ""
    stmt_date = thai_long_date(period_end_raw)
    prior_date = thai_long_date(clean(prof.get("prior_period_end")) or "")
    auditor_name = clean(prof.get("auditor_name"))
    license_no = clean(prof.get("auditor_license"))
    sign_date = clean(prof.get("sign_date"))

    status = prof.get("juristic_status", "")
    is_liq = any(k in status for k in ("เลิก", "ชำระบัญชี", "ร้าง")) or "งบเลิก" in ctx_path.parent.name
    if args.kind:
        kind = args.kind
    elif is_liq:
        kind = "cpa_liq"
    elif entity.startswith("หจก") or prof.get("auditor_type", "").startswith("TA"):
        kind = "ta"
    else:
        kind = "cpa"

    required = {"company_name": company, "period_end": stmt_date,
                "auditor_name": auditor_name, "auditor_license": license_no, "sign_date": sign_date}
    missing = [k for k, v in required.items() if not v]
    if missing:
        fail(f"unresolved required field(s) in CONTEXT (⚠ or missing): {missing}. Resolve before rendering.")

    out_dir = ctx_path.parent / "WP"
    out_dir.mkdir(parents=True, exist_ok=True)

    # ปีแรก (first-year) opinion wording: only meaningful for the standard CPA report —
    # liquidation already has its own explicit period-start token, and TA isn't in scope here.
    first_year, incorp_long = detect_first_year(
        clean(prof.get("incorporation_date")) or "",
        clean(prof.get("prior_period_end")) or "",
        period_end_raw)
    period_phrase = "ปีสิ้นสุดวันเดียวกัน"
    if first_year and incorp_long:
        period_end_era = thai_long_date_era(period_end_raw, spaced=False)
        period_phrase = (f"รอบระยะเวลาบัญชี ตั้งแต่วันที่ {incorp_long} "
                          f"(วันที่จดทะเบียนบริษัท) สิ้นสุดวันที่ {period_end_era}")

    p = {"company": company, "stmt_date": stmt_date, "prior_stmt_date": prior_date or "",
         "auditor_name": auditor_name, "license": license_no, "sign_date": sign_date,
         "city": clean(prof.get("sign_city")) or FIRM_CITY,
         "has_prior_auditor_note": prof.get("has_prior_auditor_note", ""),
         "prior_auditor_opinion": prof.get("prior_auditor_opinion", ""),
         "prior_auditor_sign_date": prof.get("prior_auditor_sign_date", ""),
         "prior_auditor_text": prof.get("prior_auditor_text", ""),
         "date_suffix": date_suffix(period_end_raw) or "",
         "period_phrase": period_phrase}

    warnings = []
    cover_kind = "ta" if kind == "ta" else "cpa"
    # Cover page carries the พ.ศ. era marker on บจ. (cpa) covers only — ground truth never
    # shows it on the TA (หจก.) cover, and the liquidation cover overrides this below anyway.
    cover_date = thai_long_date_era(period_end_raw, spaced=True) if cover_kind == "cpa" else stmt_date

    if kind == "cpa_liq":
        liq_date = args.liquidation_date
        if not liq_date:
            fail("งบเลิก (cpa_liq): --liquidation-date is required (Thai long form, e.g. "
                 "'26 ธันวาคม พ.ศ. 2568'). Ask the client / read the จดทะเบียนเลิก.")
        liq_parts = parse_thai_long_date(liq_date)
        if not liq_parts:
            fail(f"งบเลิก (cpa_liq): --liquidation-date {liq_date!r} is not a recognizable "
                 "Thai long date ('D Month [พ.ศ.] YYYY', e.g. '26 ธันวาคม พ.ศ. 2568') — "
                 "needed to build the ddmmyy filename suffix.")
        ld, lmo, ly = liq_parts
        p["date_suffix"] = f"{ld:02d}{lmo:02d}{ly % 100:02d}"
        yr = re.search(r"25\d{2}", liq_date)
        p["liq_date"] = liq_date
        p["period_start"] = args.period_start or (f"1 มกราคม พ.ศ. {yr.group(0)}" if yr else "")
        cover_date = f"{liq_date} (วันที่จดทะเบียนเลิกบริษัท)"

    cover = render_cover(cover_kind, company, cover_date, out_dir)

    if kind == "cpa":
        report = render_report_cpa(p, out_dir)
        leftover = residual_tokens(report)
        if leftover:
            fail(f"CPA report still has unfilled tokens {leftover} — this is a bug, not a data gap")
    elif kind == "cpa_liq":
        report = render_report_cpa_liquidation(p, out_dir)
        leftover = residual_tokens(report)
        if leftover:
            fail(f"งบเลิก report still has unfilled tokens {leftover} — pass --period-start if «PERIOD_START» remains")
    else:
        extra = auditor_extra(auditor_name)
        p.update({"cur_year": stmt_date.split()[-1], "prior_year": str(int(stmt_date.split()[-1]) - 1),
                  "national_id": extra.get("national_id", ""),
                  "office": args.office or extra.get("office") or FIRM_OFFICE,
                  "prior_sign_date": args.prior_sign_date or ""})
        report, flagged = render_report_ta(p, out_dir)
        for tok in flagged:
            warnings.append(f"{tok} unresolved (data not in DB/CONTEXT — flagged red in the .docx, fill by hand before use)")

    cover_left = residual_tokens(cover)
    if cover_left:
        fail(f"cover still has unfilled tokens {cover_left} — this is a bug")

    print(json.dumps({
        "ok": True, "kind": kind,
        "cover": str(cover), "report": str(report),
        "warnings": warnings,
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
