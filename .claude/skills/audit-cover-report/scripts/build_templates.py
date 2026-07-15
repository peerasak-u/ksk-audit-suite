# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx"]
# ///
"""One-time DEV script: turn ground-truth .docx into tokenized templates.

Reads real documents from 5_ตัวอย่างไฟล์ผลลัพธ์/ and writes client-neutral
template assets to this skill's assets/ dir, with every variable value replaced
by a «TOKEN». Runtime rendering (render_cover_report.py) only substitutes tokens.

Run once from the repository root when templates need regenerating:
    uv run .claude/skills/audit-cover-report/scripts/build_templates.py

Not used at runtime. The generated assets/*.docx ARE committed.

NOTE (public repo): the literal source values below (company names, auditor
names/license numbers, addresses) originally matched real client data in the
firm's private `5_ตัวอย่างไฟล์ผลลัพธ์/` folder, which is gitignored and never
committed. They have been replaced with placeholder text for this public repo.
To actually re-run this script, substitute the real values read off your own
local ground-truth documents.
"""
import re
import subprocess
import sys
from pathlib import Path

from docx import Document

ASSETS = Path(__file__).resolve().parent.parent / "assets"


def normalize_spaces(doc) -> None:
    """Collapse runs of 2+ spaces to one so rendered prose is clean. Tokens carry no
    spaces, so runtime-inserted values (e.g. the double-spaced national ID) are untouched."""
    for p in doc.paragraphs:
        if p.runs and re.search(r"  +", p.text):
            p.runs[0].text = re.sub(r"  +", " ", p.text).strip()
            for r in p.runs[1:]:
                r.text = ""


def find(tag: str, name: str) -> str:
    out = subprocess.run(
        ["find", "5_ตัวอย่างไฟล์ผลลัพธ์", "-type", "f", "-name", name],
        capture_output=True, text=True,
    ).stdout.splitlines()
    hits = [x for x in out if tag in x]
    if not hits:
        sys.exit(f"source docx not found: tag={tag} name={name}")
    return hits[0]


def retokenize(para, replacements: list[tuple[str, str]]) -> bool:
    """If the paragraph's full text contains any source value, collapse its runs
    to a single run with every value replaced by its token. Preserves first-run font."""
    full = para.text
    new = full
    for old, tok in replacements:
        new = new.replace(old, tok)
    if new == full or not para.runs:
        return False
    para.runs[0].text = new
    for r in para.runs[1:]:
        r.text = ""
    return True


def replace_whole(para, token: str) -> None:
    para.runs[0].text = token
    for r in para.runs[1:]:
        r.text = ""


def assert_clean(path: Path, banned: list[str]) -> None:
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    leaked = [b for b in banned if b in txt]
    if leaked:
        sys.exit(f"TEMPLATE LEAK in {path.name}: source value(s) not tokenized: {leaked}")


def build_cover_cpa() -> None:
    d = Document(find("S [103]", "ใบปะหน้างบการเงิน.docx"))
    repl = [("บริษัท  ตัวอย่าง เอ จำกัด", "«COMPANY»"),
            ("บริษัท ตัวอย่าง เอ จำกัด", "«COMPANY»"),
            ("31 ธันวาคม 2568", "«STMT_DATE»")]
    for p in d.paragraphs:
        retokenize(p, repl)
    out = ASSETS / "cover_cpa.docx"
    normalize_spaces(d)
    d.save(out)
    assert_clean(out, ["ตัวอย่าง เอ", "2568"])


def build_cover_ta() -> None:
    d = Document(find("S [485]", "ใบปะหน้างบการเงิน.docx"))
    repl = [("ห้างหุ้นส่วนจำกัด ตัวอย่างบี", "«COMPANY»"),
            ("31 ธันวาคม 2568", "«STMT_DATE»")]
    for p in d.paragraphs:
        retokenize(p, repl)
    out = ASSETS / "cover_ta.docx"
    normalize_spaces(d)
    d.save(out)
    assert_clean(out, ["ตัวอย่างบี", "2568"])


def build_report_cpa() -> None:
    """From client [233] — a CPA report WITH the optional prior-auditor block,
    so one template serves both (block removed when not needed)."""
    d = Document(find("S [233]", "2 หน้ารายงาน บริษัท ตัวอย่างซี จำกัด 311268.docx"))
    repl = [("บริษัท ตัวอย่างซี จำกัด", "«COMPANY»"),
            ("31 ธันวาคม 2568", "«STMT_DATE»"),
            ("31 ธันวาคม 2567", "«PRIOR_STMT_DATE»"),
            ("นางสาวตัวอย่าง ผู้สอบเอ", "«AUDITOR_NAME»"),
            ("00001", "«LICENSE»"),
            ("20 เมษายน 2569", "«SIGN_DATE»")]
    for p in d.paragraphs:
        t = p.text.strip()
        if t == "เรื่องอื่นๆ":
            replace_whole(p, "«IF_PRIOR»เรื่องอื่นๆ")
        elif t.startswith("งบการเงินของ") and "ตรวจสอบโดยผู้สอบบัญชีอื่น" in t:
            replace_whole(p, "«PRIOR_AUDITOR_TEXT»")
        elif t == "ขอนแก่น":
            replace_whole(p, "«CITY»")
        else:
            retokenize(p, repl)
    out = ASSETS / "report_cpa.docx"
    normalize_spaces(d)
    d.save(out)
    assert_clean(out, ["ตัวอย่างซี", "ผู้สอบเอ", "00001", "2568", "2567"])


def build_report_ta() -> None:
    """From client [485] — the RD-style TA report. Signature block needs the
    auditor's national ID and office address (from Database ผู้สอบ)."""
    d = Document(find("S [485]", "2 รายงานการตรวจสอบและรับรองบัญชี หจก ตัวอย่างบี 311268.docx"))
    repl = [("ห้างหุ้นส่วนจำกัด     ตัวอย่างบี", "«COMPANY»"),
            ("ห้างหุ้นส่วนจำกัด   ตัวอย่างบี", "«COMPANY»"),
            ("ห้างหุ้นส่วนจำกัด ตัวอย่างบี", "«COMPANY»"),
            ("31 ธันวาคม 2568", "«STMT_DATE»"),
            ("ปี 2568", "ปี «CUR_YEAR»"),
            ("ปี 2567", "ปี «PRIOR_YEAR»"),
            ("2567 (ปีก่อน)", "«PRIOR_YEAR» (ปีก่อน)"),
            ("นางตัวอย่าง  ผู้สอบบี", "«AUDITOR_NAME»"),
            ("นางตัวอย่าง ผู้สอบบี", "«AUDITOR_NAME»"),
            ("00002", "«LICENSE»"),
            ("1  2  3  4  5  6  7  8  9  0  1  2  3", "«NATIONAL_ID»"),
            ("0/0 ถ.ตัวอย่าง ต.ตัวอย่าง อ.ตัวอย่าง จ.ตัวอย่าง", "«OFFICE»"),
            ("28 พฤษภาคม2568", "«PRIOR_SIGN_DATE»"),
            ("16 พฤษภาคม 2569", "«SIGN_DATE»")]
    for p in d.paragraphs:
        retokenize(p, repl)
    out = ASSETS / "report_ta.docx"
    normalize_spaces(d)
    d.save(out)
    assert_clean(out, ["ตัวอย่างบี", "ผู้สอบบี", "00002", "0/0 ถ.ตัวอย่าง"])


def build_report_cpa_liquidation() -> None:
    """From client [404] — a CPA report for a dissolved company (งบเลิก). Differs from the
    standard CPA report: addressed to กรรมการ, adds an emphasis-of-matter paragraph, and
    uses ผู้ชำระบัญชี throughout. Dates are the liquidation date + the period start."""
    d = Document(find("S [404]", "2 หน้ารายงานผู้สอบ บริษัท ตัวอย่างดี จำกัด 261268.docx"))
    repl = [("บริษัท ตัวอย่างดี จำกัด", "«COMPANY»"),
            ("26 ธันวาคม พ.ศ. 2568", "«LIQ_DATE»"),
            ("1 มกราคม  พ.ศ. 2568", "«PERIOD_START»"),
            ("1 มกราคม พ.ศ. 2568", "«PERIOD_START»"),
            ("นางสาวตัวอย่าง ผู้สอบเอ", "«AUDITOR_NAME»"),
            ("00001", "«LICENSE»"),
            ("16 กุมภาพันธ์ 2569", "«SIGN_DATE»")]
    for p in d.paragraphs:
        if p.text.strip() == "ขอนแก่น":
            replace_whole(p, "«CITY»")
        else:
            retokenize(p, repl)
    out = ASSETS / "report_cpa_liquidation.docx"
    normalize_spaces(d)
    d.save(out)
    assert_clean(out, ["ตัวอย่างดี", "ผู้สอบเอ", "00001", "2568"])


def main() -> None:
    if not Path("5_ตัวอย่างไฟล์ผลลัพธ์").is_dir():
        sys.exit("run from repository root (5_ตัวอย่างไฟล์ผลลัพธ์/ not found)")
    ASSETS.mkdir(exist_ok=True)
    build_cover_cpa()
    build_cover_ta()
    build_report_cpa()
    build_report_cpa_liquidation()
    build_report_ta()
    print("built templates:", ", ".join(p.name for p in sorted(ASSETS.glob("*.docx"))))


if __name__ == "__main__":
    main()
